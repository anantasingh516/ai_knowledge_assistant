import streamlit as st
import requests
import io
import os
import json
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from datetime import datetime

st.set_page_config(
    page_title="AI Knowledge Assistant",
    page_icon="🤖",
    layout="wide"
)

API_URL = "http://127.0.0.1:8000/ask"
LOG_FILE_PATH = "logs/query_history.jsonl"
EVAL_SUMMARY_FILE = "logs/eval_metrics_summary.json"

if "messages" not in st.session_state:
    st.session_state.messages = []
if "last_response" not in st.session_state:
    st.session_state.last_response = None

def generate_word_doc(question, answer, citations):
    """Generates a styled Word document for export."""
    doc = Document()
    doc.add_heading("AI Knowledge Assistant - Export Report", level=1)
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
    
    doc.add_heading("User Question", level=2)
    doc.add_paragraph(question)
    
    doc.add_heading("Grounded AI Answer", level=2)
    doc.add_paragraph(answer)
    
    doc.add_heading("Source Citations", level=2)
    for c in citations:
        doc.add_paragraph(f"• Source: {c['source']} | ID: {c['id']} | Match Distance: {c['match_score']}")
        
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def generate_excel_sheet(question, answer, citations):
    """Generates a structured Excel spreadsheet for export."""
    data = []
    if citations:
        for c in citations:
            data.append({
                "Question": question,
                "Answer": answer,
                "Citation Source": c['source'],
                "Chunk ID": c['id'],
                "Match Score": c['match_score']
            })
    else:
        data.append({"Question": question, "Answer": answer, "Citation Source": "N/A", "Chunk ID": "N/A", "Match Score": "N/A"})
        
    df = pd.DataFrame(data)
    bio = io.BytesIO()
    with pd.ExcelWriter(bio, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name="RAG Report")
    bio.seek(0)
    return bio
with st.sidebar:
    st.title("⚙️ Control Center")
    st.markdown("---")
    st.subheader("📁 Upload New Documents")
    uploaded_file = st.file_uploader(
        "Drop a file to expand the AI's knowledge base:", 
        type=["txt", "pdf", "docx", "csv"]
    )
    
    if uploaded_file is not None:
        if st.button("🚀 Process & Ingest File", use_container_width=True):
            with st.spinner("Streaming data bytes to core embedding engine..."):
                try:
                    files = {"file": (uploaded_file.name, uploaded_file.getvalue(), uploaded_file.type)}
                    upload_url = "http://127.0.0.1:8000/upload"
                    
                    response = requests.post(upload_url, files=files, timeout=None)
                    
                    if response.status_code == 200:
                        st.success(f"Success! {uploaded_file.name} is now live in the knowledge matrix.")
                    else:
                        st.error(f"Error packing document: {response.json().get('detail')}")
                except Exception as e:
                    st.error(f"Connection to backend failed: {e}")

    st.markdown("---")
    st.subheader("Database Management")
    if st.button("🔄 Trigger Manual Re-Index", use_container_width=True, type="primary"):
        with st.spinner("Scanning data vault and updates..."):
            try:
                from core.search_engine import VectorSearchEngine
                engine = VectorSearchEngine()
                engine.index_processed_vault()
                st.success("Vector Database Re-Indexed Successfully!")
            except Exception as e:
                st.error(f"Re-indexing failed: {e}")
                
    st.markdown("---")
    st.caption("AI Assistant Version 1.2 (Evaluation & Telemetry Live)")
tab1, tab2 = st.tabs(["💬 Assistant Chatbot", "📊 System Performance & Evaluation Dashboard"])

with tab1:
    st.subheader("Interactive Knowledge Stream")
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])
            
    if user_query := st.chat_input("Ask a question about internal procedures or algorithms..."):
        st.session_state.messages.append({"role": "user", "content": user_query})
        with st.chat_message("user"):
            st.markdown(user_query)
            
        with st.chat_message("assistant"):
            response_placeholder = st.empty()
            
            with st.spinner("Searching document vault and generating grounded response..."):
                try:
                    payload = {"question": user_query, "top_k": 2}
                    res = requests.post(API_URL, json=payload, timeout=None)
                    
                    if res.status_code == 200:
                        data = res.json()
                        ai_answer = data["answer"]
                        citations = data["citations"]
                        
                        st.session_state.last_response = {
                            "question": user_query,
                            "answer": ai_answer,
                            "citations": citations
                        }
                        
                        response_placeholder.markdown(ai_answer)
                        st.session_state.messages.append({"role": "assistant", "content": ai_answer})
                        
                        if citations:
                            with st.expander("📚 View Grounded Source References"):
                                for i, cite in enumerate(citations):
                                    st.markdown(f"**Reference {i+1}:** {cite['source']}")
                                    st.caption(f"Chunk ID: `{cite['id']}` | Semantic Proximity Distance: `{cite['match_score']}`")
                    else:
                        response_placeholder.error(f"Backend Server Error (Code {res.status_code})")
                except requests.exceptions.ConnectionError:
                    response_placeholder.error("Could not connect to backend server. Make sure `python run.py` is active on port 8000!")

    
    if st.session_state.last_response:
        st.markdown("---")
        st.subheader("Export Last Generated Response")
        
        q = st.session_state.last_response["question"]
        a = st.session_state.last_response["answer"]
        c = st.session_state.last_response["citations"]
        
        col1, col2 = st.columns(2)
        with col1:
            doc_bytes = generate_word_doc(q, a, c)
            st.download_button(
                label="📄 Download Report as Word (.docx)",
                data=doc_bytes,
                file_name=f"rag_report_{datetime.now().strftime('%M%S')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        with col2:
            xls_bytes = generate_excel_sheet(q, a, c)
            st.download_button(
                label="📊 Download Table Data as Excel (.xlsx)",
                data=xls_bytes,
                file_name=f"rag_report_{datetime.now().strftime('%M%S')}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True
            )

with tab2:
    st.subheader("🔍 Operational Telemetry & Evaluation Metrics")
    if os.path.exists(LOG_FILE_PATH):
        st.markdown("### 🪵 Query Operations Real-time Log Stream")
        
        log_data = []
        with open(LOG_FILE_PATH, "r", encoding="utf-8") as f:
            for line in f:
                if line.strip():
                    log_data.append(json.loads(line))
        df = pd.DataFrame(log_data)
        display_df = df[["timestamp", "query", "latency_seconds", "context_word_overlap_ratio"]].tail(10)
        st.dataframe(display_df, use_container_width=True)
        st.markdown("### 📈 Analytical System Operational Trends")
        g_col1, g_col2 = st.columns(2)
        
        with g_col1:
            fig1, ax1 = plt.subplots(figsize=(6, 3))
            ax1.plot(df["latency_seconds"].tail(15), marker='o', color='#1E88E5', linewidth=2)
            ax1.set_title("Ollama LLM Processing Latency (Seconds)", fontsize=10)
            ax1.set_ylabel("Seconds")
            ax1.grid(True, linestyle='--', alpha=0.5)
            st.pyplot(fig1)
            
        with g_col2:
            fig2, ax2 = plt.subplots(figsize=(6, 3))
            ax2.plot(df["context_word_overlap_ratio"].tail(15), marker='s', color='#2E7D32', linewidth=2)
            ax2.set_title("Answer Context Grounding Ratio (Lexical Overlap)", fontsize=10)
            ax2.set_ylabel("Overlap Percentage")
            ax2.grid(True, linestyle='--', alpha=0.5)
            st.pyplot(fig2)
    else:
        st.info("💡 Telemetry log lines are empty. Submit your first query in Chat to start generating analytical charts.")

    
    st.markdown("---")
    st.markdown("### 🎯 Automated Validation Suite (Ground Truth Benchmarks)")
    
    c1, c2 = st.columns([1, 2])
    with c1:
        if st.button("🚀 Run Evaluation Suite Benchmark", use_container_width=True):
            with st.spinner("Querying vector indexing coordinates programmatically..."):
                try:
                    from core.evaluator import run_automated_evaluation
                    run_automated_evaluation(k=3)
                    st.success("RAG Metrics recalculated successfully!")
                except Exception as e:
                    st.error(f"Execution error running automated benchmarking system: {e}")
    if os.path.exists(EVAL_SUMMARY_FILE):
        with open(EVAL_SUMMARY_FILE, "r", encoding="utf-8") as eval_f:
            metrics = json.load(eval_f)
            
        st.markdown("#### Baseline Target Statistics (Current Calculation Run)")
        m_col1, m_col2, m_col3 = st.columns(3)
        m_col1.metric("Recall @ k Accuracy", f"{int(metrics['avg_recall_at_k'] * 100)}%")
        m_col2.metric("Citation Metadata Coverage", f"{int(metrics['avg_citation_coverage'] * 100)}%")
        m_col3.metric("Retrieval Search Horizon Depth", f"k = {metrics['evaluated_at_k']}")
        st.markdown("**Test Set Coverage Log:**")
        detailed_df = pd.DataFrame(metrics["detailed_breakdown"])
        st.table(detailed_df)
    else:
        st.warning("⚠️ No system verification results found on disk. Click the button above to execute the initial baseline calculation.")